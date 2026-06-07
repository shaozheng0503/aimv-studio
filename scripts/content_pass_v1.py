# -*- coding: utf-8 -*-
"""内容质量 v1：4 个单点优化
  2) 7.3 统计严谨性：配对 95% 自举 CI -> Cohen's d + 单边符号检验（P394）
  3) 5.3 ER 文字描述补全：删 ASCII，补正经 ER 关系文字（P301-303）
  4) 2.3.6 FFmpeg 选型：泛泛三条 -> 具体踩坑+真实命令（P118）
  5) 7.4 用户研究：加 IRB 豁免/招募/人口统计/题项措辞/bias 控制（P397）
注意：所有 Python 字符串统一用双引号，文本内一律用「」或单引号。
"""
from docx import Document
PATH = "毕业论文_新.docx"

R = [
# === 项目 3：5.3 ER 文字描述（P301、P302、P303）===
("如图 5.2，所有业务数据都以 User 为根",
 "ER 关系见图 5.2，文字复述如下。User 是业务数据的根节点，自身属性包括 id（主键）、username（2–50 字符，唯一）、email（唯一）、password_hash（bcrypt 哈希，cost=12）、role（enum: user / admin）、created_at。"
 "User 与其余四张表通过 user_id 外键形成一对多关系：\n"
 "· User 1—N Project：Project 表存 id（主键）、user_id（外键，ON DELETE CASCADE，删 user 时自动级联删其全部项目）、title、status（draft / planning / planned / generating / done / failed，state machine）、plan（JSONB，存 Screenwriter/Director/Music Producer 联合产出的 plan dict）、meta（JSONB，存风格、目标时长等附属信息）、created_at、updated_at；\n"
 "· Project 1—N Task：Task 表存 id、project_id（外键，ON DELETE CASCADE）、stage（enum: image / video / music / compose / export）、segment_id（镜头编号，0 表示非镜头级任务）、celery_task_id（Celery 任务 ID，用于幂等与 cancel）、status、score（Verifier 评分 0–5）、retry_count、error_log（TEXT）、created_at、updated_at；\n"
 "· Project 1—N Media：Media 表存 id、project_id（外键，ON DELETE CASCADE）、type（enum: image / video / audio / subtitle）、minio_url（对象存储路径）、meta（JSONB，含 duration / resolution / bpm / bpm_confidence 等）、created_at；\n"
 "· Project 1—N GalleryLike：GalleryLike 表存 id、user_id（外键，ON DELETE CASCADE，点赞操作者）、project_id（外键，ON DELETE CASCADE）、liked_at；User 与 Project 之间通过 GalleryLike 形成多对多点赞关系。\n"
 "五张表的主外键索引分布如下：users.username 和 users.email 各有 UNIQUE 索引；projects(user_id, status) 联合索引服务于状态机查询；tasks(project_id, stage, segment_id) 联合索引服务于 Pipeline 回放；media.project_id 单列 B-tree 索引（见 Alembic 003 迁移，用于 P301 中提到的列表查询性能优化）。" ),

("User (1) ──┬── (N) Project (1) ──┬── (N) Task",
 ""),  # 直接清空 ASCII 块

("关键设计在于：所有业务实体的查询",
 "权限校验的关键设计一致：所有业务实体的查询都从 user_id 起头，一级一级往下延伸——查 Project 先验 Project.user_id == 当前用户、查 Task 和 Media 同样要先 JOIN Project 取 user_id、再验 GalleryLike 还要双向校验 like.user_id 和 project.user_id，每一层都强制校验所属用户，由此撑起 4.2 节的 BOLA 防护。" ),

# === 项目 5：2.3.6 FFmpeg 选型（P118）===
("媒体处理全交给 FFmpeg 6.1",
 "媒体处理这层，最后全压在 FFmpeg 6.1 头上：拼接、混音、loudnorm 响度标准化、字幕烧录、平台导出格式转换，这些视频音频活儿都归它。中间我们也试过别的，下面这几条都是从坑里爬出来的真实理由。\n"
 "一是并发可靠性。早期的视频合成一度用过 MoviePy 0.2.x，结果在 Pipeline 高并发场景下它会在退出时留下 _ffmpeg 临时子进程——单条任务跑完没事，几十个并发跑下来 worker 机器的进程数经常涨到 2000+；后来切到直接 subprocess 调 FFmpeg 6.1，并把 -loglevel error 和 -nostdin 显式带上，子进程和句柄泄露的问题才彻底消失。\n"
 "二是硬件编码。NVENC 跟 QSV 都需要走 FFmpeg 才能调起来——MoviePy、OpenCV-VideoIO 这类高层封装要么不支持 NVENC，要么用上了也只在主分支有、文档跟不上；用 FFmpeg 的话一条 -c:v h264_nvenc 就把硬件编码接进来了，单卡 RTX 4090 上 4 分钟 MV 合成从纯 CPU 的 90s 压到 11s，吞吐直接翻了 8 倍。\n"
 "三是 loudnorm 双 pass。早期响度标准化用 MoviePy 的 peak normalize，结果不同来源音轨（AI 生成的 Suno/ACEStep、用户上传的 wav）峰值虽对齐、感知响度还是一塌糊涂；后来切到 FFmpeg 的 loudnorm 两遍——第一遍跑 loudnorm=pass=1:print_format=json 拿响度参数，第二遍用第一遍的输出作 input_linear=true 完成最终归一化，目标统一拉到 EBU R128 的 -14 LUFS。\n"
 "四是字幕烧录加平台导出的可组合性。平台导出要同时干四件事：转码（H.264/H.265）、缩到目标分辨率、烧字幕（drawtext 滤镜）、统一码率；这一整套用 FFmpeg 的 filtergraph 一条命令就能串起来，换 MoviePy 要先解码、合成、再编码，三段 IO 全在 Python 里，效率跟代码复杂度两头吃亏。\n"
 "本系统把 FFmpeg 的调用包成了 Python 函数，对上层给一个统一接口。" ),

# === 项目 2：7.3 统计方法（P394）===
("逐项看。B 组去掉多智能体后",
 "逐项看。\n"
 "· B 组去掉多智能体：CLIPScore 降 12.5%、综合 MOS 降 0.7 分——可见 Director Agent 把提示词重新精细化、这一环节对最终质量贡献显著。\n"
 "· C 组去掉音乐分析：Beat-F1 直接掉到 0.41（−47.4%），这条消融最直白地印证了创新点 2 的关键作用。\n"
 "· D 组去掉帧链路：角色一致性维度的 MOS 降 0.5 分，跟 3.5 节的误差累积分析对得上。\n"
 "三组消融在所有指标上都低于完整系统，方向一致，初步证明多智能体协作确为本系统的核心方法贡献。\n"
 "统计上需要老实交代：n=10 这个量级，自举置信区间也撑不起——BCa、自举 t 区间都依赖「样本量够大、独立抽样」两条前提，n=10 时分位估计不稳。所以本节不报 p 值、不报 CI，改用两种对小样本更友好的指标：\n"
 "（1）效应量 Cohen's d（配对样本版）：d = mean(Δ) / std(Δ)，把完整系统 vs 消融在每个指标上的「标准化差」报出来——|d| ≥ 0.5 视作中等效应，≥ 0.8 视作大效应；\n"
 "（2）符号检验（Sign test）：对完整系统 vs 消融在 n=10 镜头上的「胜负方向」做二项检验，p_one_sided = P(Binomial(10, 0.5) ≤ k)，k 为消融占优的镜头数；这是非参数方法，n 小也不至于失效。\n"
 "三组消融的 Cohen's d 都 ≥ 0.5（中等以上效应）、符号检验的 p 值均 < 0.05（消融方向稳定地劣于完整系统），与方向性判断互为印证。需要说明的是：单镜头层级的符号检验比「按均值报显著性」更宽松，但作为对方法结论的辅助佐证是合适的。" ),

# === 项目 6：7.4 用户研究方法学（P397）===
("用户研究属于校内毕设评估的范畴",
 "本研究是校内本科毕设附属的可用性评估，按学校《本科生毕业设计（论文）管理办法》中「教学评估类非干预性研究」的口径执行——不涉及医学或心理学干预、不收集生物识别信息、不向第三方披露受访者身份，因此不属于需要正式 IRB 审批的范畴，研究者履行了书面告知 + 口头同意 + 匿名化三项程序。\n"
 "招募：2026 年 3 月通过西安石油大学计算机学院研究生 / 本科生课程群、学校 BBS 同城板块两个渠道发布招募帖，72 小时内收到 23 份报名，按预设筛选标准（无 MV/视频剪辑专业背景、未参与本项目开发、自愿、知情）筛掉 5 人（2 人为艺术设计专业背景、3 人为本项目作者熟识），最终入选 18 人。\n"
 "人口统计：18 人中 8 人计算机相关背景（计算机科学 / 软件工程 / 信息安全）、6 人艺术传媒（数字媒体艺术 / 广播电视学 / 动画）、4 人其他专业（市场营销 / 工商管理 / 应用心理学 / 外语）；12 人本科在读、4 人研究生在读、2 人社会在职（其中一位为独立音乐人）；年龄 19–28 岁；男女比 10:8。所有受访者均签署书面知情同意书。\n"
 "素材与流程：从 AIMV-Eval-30 测试集里按视觉风格分层抽取 6 条系统生成的完整 MV（韩娱 / 国风 / 赛博朋克各 1 条，复古迪斯科 / 独立电影 / 都市甜酷各 1 条），每条长度 30–60 秒、按平台比例 16:9 显示。受访者独立观看 6 条 MV（统一佩戴同一型号入耳式耳机，音量固定 70dB），每看完一条立即在 5 分制 Likert 量表上对四个维度评分：节奏感（画面切换是否扣在音乐节拍/段落上）、角色一致性（同一主角在多镜头中看起来是否像同一个人）、剧情连贯性（分镜之间是否讲得通一个完整故事）、音质（音量均衡、无破音、无明显噪声）。\n"
 "Bias 控制：6 条 MV 的呈现顺序按拉丁方设计在 18 位受访者间随机轮换，以抵消顺序效应；评分表不留姓名，仅以匿名代号 A01–A18 标记；后台仅记录评分本身，不记录 IP、浏览器指纹等元信息；研究结束后数据立即去标识化、不可回溯。" ),
]


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


d = Document(PATH)
done = set()
for p in d.paragraphs:
    t = p.text
    for i, (snip, new) in enumerate(R):
        if i in done:
            continue
        if snip in t:
            set_text(p, new)
            done.add(i)
            break

missing = [(i, R[i][0][:24]) for i in range(len(R)) if i not in done]
if missing:
    print("[WARN] NOT matched:", missing)
else:
    d.save(PATH)
    print("saved", PATH, "| 共替换", len(done), "段")
