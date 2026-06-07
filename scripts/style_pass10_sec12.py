# -*- coding: utf-8 -*-
"""pass10 续：改写 1.2 国内外研究发展现状（6 段）。保留全部引用标记。"""
from docx import Document
PATH = "毕业论文_新.docx"

REPLACEMENTS = [
    ("自 DDPM 提出以来",
     "图像和视频生成这几年基本被扩散模型主导。从 DDPM [3, 4] 起步，到 Stable Diffusion 3 把 Rectified Flow 接进 "
     "Transformer、把保真度又推高一截 [21]，这条路线已相当成熟。视频侧更热闹：OpenAI 的 Sora 走的是基于 DiT 的"
     "“视频世界模型”路子 [22]，Google DeepMind 的 Veo 系列把单段时长做到了分钟级，阿里通义 2025 年开源的 "
     "Wan2.2 14B（Apache 2.0）则把开源视频模型第一次顶到了准商业水准。音频这边也没落下，MusicGen [6] 给出了可控生成"
     "的框架，Suno、Udio 这类产品甚至能从歌词和风格直接出一整首歌——只是它们都只给音频，吐不出能二次剪辑的画面素材。"
     "所以单步生成的质量其实早够用了，缺的一直是把多种模态、多个镜头端到端串起来协同的那一环。"),

    ("2023 年以后，多智能体协作",
     "多智能体协作大概从 2023 年起火起来 [16, 19]，几项工作奠定了基调。AutoGen [12] 最早把“多个 LLM 像人一样对话协作”"
     "这件事跑通；MetaGPT [13] 干脆把一个软件公司里的不同岗位拆成各自独立的 Agent，实验也确认了——把角色分清楚，"
     "复杂任务完成得更好；Multi-Agent Debate [15] 让几个模型互相辩，单模型那些张口就来的事实错误明显少了；"
     "ChatDev [14] 则用瀑布式的多 Agent 把软件从需求一路做到代码。反思这条线上，Reflexion [17] 很关键，"
     "本系统里 Verifier 那套“打分不行就重来”的循环，灵感正是从它来的。"),

    ("综合若干代表性综述",
     "几篇综述 [16, 19] 把多 Agent 相比单 Agent 的好处大致归到了几种场景：决策本身很杂、要同时权衡好几类异质目标的；"
     "需要“先生成再自我批评”这种反思闭环的；还有靠“扮演某个角色”就能把模型能力激发出来的。MV 创作恰好把这几样占全了"
     "——它既要同时定下剧本、画面和音乐，又离不开质检返工，每个 Agent 还都得入戏。这也是本课题一上来就选多智能体当主方法的原因。"),

    ("首先考察国外产品",
     "先看国外。Runway Gen-3、Pika 1.5 大体还是“一个镜头出一个结果”；Suno 的 MV Mode 虽然能“给歌配画面”，画质却比较一般。"
     "国内这边，可灵、即梦、Vidu、智谱清影在视频生成上都不弱，但还没有哪家专门把“给一首歌、自动剪成 MV”这条完整链路做下来。"
     "为了看清差距，表 1.1 从六个维度把国内外这些产品和本系统摆在一起比。"),

    ("由表 1.1 可见一个共同点",
     "比下来有个很一致的现象：绝大多数产品的力气都花在“单步生成”这一层，说到底就是把一句提示词变成一张图或一段视频。"
     "连覆盖了多模态的 Suno MV Mode，跨镜头一致性也只靠“画风别差太多”这种弱约束撑着，音乐对齐更是停在 BPM，"
     "谈不上副歌、主歌这种段落级的语义对齐。AIMV Studio 不太一样的地方在于，多模态覆盖、跨镜头一致性、音乐对齐、"
     "自动化、多平台分发这五件事是一起做的；整张表里，真正照着“从一句创意到能直接发布的成片”这条完整链路去设计的，就这一个。"),

    ("这种差别根本上源于方法论",
     "差别的根子在方法论。别人多半是“一条 Prompt 直接喂给一个大模型”的单步打法，AIMV Studio 走的是"
     "“多 Agent 拆决策 + 模型路由 + 末帧链路”的组合拳。换句话说，它不再死磕单个模型的极限，"
     "而是把 MV 生成当成一件需要逐层规划的系统工程来做。"),
]


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


d = Document(PATH)
done = set()
for p in d.paragraphs:
    t = p.text.strip()
    for i, (prefix, new) in enumerate(REPLACEMENTS):
        if i in done:
            continue
        if t.startswith(prefix):
            set_text(p, new)
            done.add(i)
            print(f"[ok] #{i}: {prefix[:16]}…")
            break

missing = [REPLACEMENTS[i][0][:16] for i in range(len(REPLACEMENTS)) if i not in done]
if missing:
    print("[WARN] not matched:", missing)
else:
    d.save(PATH)
    print("saved", PATH)
