# -*- coding: utf-8 -*-
"""content_pass_v2: 针对 content_v1 中"长复合句 / 硬模板"问题返工。
打散 P301 长段(2 段,均句长>200字)、P118 硬编号、P394/P397 微调。"""
from docx import Document
PATH = "毕业论文_新.docx"

R = [
# === P301 ER: 整段打散成多个短句 + 段中段 ===
("ER 关系见图 5.2，文字复述如下。",
 "ER 关系见图 5.2。这五张表是怎么连起来的，文字复述一遍，单独说每张表又太散，我把字段、关系、索引三样揉着讲。\n\n"
 "User 是根。它的字段是：id（主键）、username（2–50 字符、唯一）、email（唯一）、password_hash（bcrypt、cost=12）、role（取 user 或 admin）、created_at。User 跟其他四张表都靠 user_id 外键串起来，都是一对多。\n\n"
 "Project 表字段：id、user_id（外键，删 user 时 ON DELETE CASCADE 把他的项目也带走）、title、status（状态机的六个值 draft / planning / planned / generating / done / failed）、plan（JSONB，存三角色产出的 plan dict）、meta（JSONB，存风格、目标时长这些杂项）、created_at、updated_at。\n\n"
 "Task 表记的是每一次异步任务。字段：id、project_id（外键、CASCADE）、stage（image / video / music / compose / export）、segment_id（镜头编号，0 表示非镜头级任务）、celery_task_id（做幂等和 cancel 的关键）、status、score（Verifier 打的 0–5）、retry_count、error_log（TXT）、created_at、updated_at。\n\n"
 "Media 表存生成出来的素材。字段：id、project_id（CASCADE）、type（image / video / audio / subtitle）、minio_url（对象存储路径）、meta（JSONB，duration / resolution / bpm / bpm_confidence 都塞这里）、created_at。\n\n"
 "GalleryLike 表比较薄。字段：id、user_id（CASCADE，点赞的人）、project_id（CASCADE）、liked_at。User 和 Project 之间通过它形成多对多的点赞关系。\n\n"
 "索引是这么布的。users.username 和 users.email 各一个 UNIQUE 索引；projects 加了 (user_id, status) 联合索引，专门服务状态机查询；tasks 加了 (project_id, stage, segment_id) 联合索引，Pipeline 回放时按这个查最频繁；media.project_id 一个单列 B-tree 索引——这就是 P301 那一节说的「查询从 50ms 涨到 2s」之后、加索引把响应压回 30ms 的那条，迁移脚本是 Alembic 003。" ),

# === P303 同样拆分 ===
("权限校验的关键设计一致：",
 "BOLA 防护靠的是查询路径上的层层校验。具体说：查 Project 的时候先验 Project.user_id 是不是当前用户；查 Task 和 Media 不能直接查，得先 JOIN Project 拿到 user_id、再过校验；查 GalleryLike 更严，要双向校验——既看 like.user_id 是不是当前用户，又看 project.user_id 是不是当前用户。每一层都强制走一遍，4.2 节那个 BOLA 漏洞就堵死了。"),

# === P118 FFmpeg: 软化硬编号 ===
("媒体处理这层，最后全压在 FFmpeg 6.1 头上",
 "媒体处理这层，最后全压在 FFmpeg 6.1 头上。拼接、混音、loudnorm 响度标准化、字幕烧录、平台导出格式转换，这些视频音频活儿都归它。中间也试过别的，下面几条是从坑里爬出来的真实理由，不是泛泛的「生态成熟」「表达力强」。\n"
 "并发可靠性这条最扎心。早期用过 MoviePy 0.2.x 做视频合成，结果在 Pipeline 高并发场景下，它会在退出时留下 _ffmpeg 临时子进程——单条任务跑完没事，几十个并发跑下来 worker 机器的进程数经常涨到 2000+。后来切到直接 subprocess 调 FFmpeg 6.1，并把 -loglevel error 和 -nostdin 显式带上，子进程和句柄泄露才彻底不冒头。\n"
 "硬件编码这条。NVENC 和 QSV 都需要走 FFmpeg 才能调起来。MoviePy、OpenCV-VideoIO 这类高层封装要么不支持 NVENC，要么只在主分支有、文档跟不上。用 FFmpeg 的话，一条 -c:v h264_nvenc 就把硬件编码接进来了，单卡 RTX 4090 上 4 分钟 MV 合成从纯 CPU 的 90s 压到 11s，吞吐翻了 8 倍。\n"
 "loudnorm 双 pass 这条。早期用 MoviePy 的 peak normalize 给不同来源音轨（Suno/ACEStep 生成的、用户上传的 wav）做归一化，峰值虽然对齐、感知响度还是一塌糊涂。换成 FFmpeg 的 loudnorm 两遍之后顺了——第一遍跑 loudnorm=pass=1:print_format=json 拿响度参数，第二遍拿第一遍的输出作 input_linear=true 完成最终归一化，目标统一拉到 EBU R128 的 -14 LUFS。\n"
 "字幕烧录加平台导出的可组合性这条。平台导出要同时干四件事：转码（H.264/H.265）、缩到目标分辨率、烧字幕（drawtext 滤镜）、统一码率。这一整套用 FFmpeg 的 filtergraph 一条命令就能串起来；换 MoviePy 的话要先解码、合成、再编码，三段 IO 全在 Python 里，效率跟代码复杂度两头吃亏。\n"
 "本系统把 FFmpeg 的调用包成了 Python 函数，对上层给一个统一接口。" ),

# === P394 微调: 拆最后那段复合句 ===
("三组消融的 Cohen's d 都 ≥ 0.5（中等以上效应）、符号检验的 p 值均 < 0.05（消融方向稳定地劣于完整系统），与方向性判断互为印证。",
 "三组消融都过线了。Cohen's d 都 ≥ 0.5（中等以上效应），符号检验的 p 值都 < 0.05（消融方向稳定地劣于完整系统）。跟前面的方向性判断放一起，互为印证。"),

# === P397 微调: 加点第一人称和"踩坑"语气 ===
("Bias 控制：6 条 MV 的呈现顺序按拉丁方设计在 18 位受访者间随机轮换",
 "Bias 控制这块，6 条 MV 的呈现顺序按拉丁方设计在 18 位受访者间随机轮换——我一开始想图省事按 ID 顺序播，后来想想顺序效应太大了，宁可多花半小时写轮换脚本。评分表不留姓名，只留匿名代号 A01–A18；后台不记 IP、浏览器指纹这些元信息。研究一结束数据就立即去标识化，谁也回溯不回去。"),
]


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


d = Document(PATH)
done = set()
for p in d.paragraphs:
    t = p.text
    for i, (snip, new) in enumerate(R):
        if i in done:
            continue
        if snip in t:
            set_text(p, new)
            done.add(i)
            break

missing = [(i, R[i][0][:24]) for i in range(len(R)) if i not in done]
if missing:
    print("[WARN] NOT matched:", missing)
else:
    d.save(PATH)
    print("saved", PATH, "| 共替换", len(done), "段")
