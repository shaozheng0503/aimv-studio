# -*- coding: utf-8 -*-
"""pass10 续：改写 3.4 多智能体规划链路（37段）。逐字保留全部参数/字段/引用/型号。"""
from docx import Document
PATH = "毕业论文_新.docx"

R = [
("继而逐一展开四个 Agent 的关键属性",
 "多智能体规划链路是本系统方法层面的核心。这一节这么走：先讲 Agent 的角色和协作要素，顺带说清 State 字段怎么读写、输出走什么协议（3.4.1）；"
 "再讲四个角色为什么这么拆（3.4.2）；接着把四个 Agent 的关键属性一个个铺开（3.4.3）；最后落到 LangGraph 编排、协作流转，以及跟单 Agent 的对比（3.4.4）。整体协作流转见图 3.2。"),

("Verifier 单独审查每个生成的镜头，分数不达标即按四档逐级回退",
 "图 3.2 画的是一套“串行规划 + 反思分支”的双轨编排。上面那条轨，Screenwriter、Director、Music Producer 按严格的语义依赖顺序一个接一个串下来；"
 "下面那条轨，Verifier 单独审查每个生成的镜头，分数不达标就按四档逐级回退。上下文都由 LangGraph 共享的 State 扛着，而且是字段级的"
 "——每个 Agent 只读自己要用的那几个字段，长上下文带来的职责漂移也就被压住了。"),

("归纳为五个要素",
 "本文把一个 LLM 智能体的工程构成拆成五样：角色定位、输出目标、可用工具集、底层 LLM 选型、超参数配置。角色定位划定它在系统里管哪一摊；"
 "输出目标规定它必须交出什么样的结构化结果；可用工具集决定它能不能调外部能力（文件读写、网络检索、媒体生成 API 之类）；"
 "底层 LLM 选型决定它能想多深、上下文能装多少；超参数那一组（temperature、max_tokens、top_p）则在创造性和稳定性之间拿捏。"
 "至于 AIMV Studio 给四个 Agent 各自配了什么角色、目标、底层模型和温度，留到 3.4.3 节细讲。"),

("常见的组织方式可归为以下四类",
 "几个 Agent 得按某种方式组织起来，才能合力把复杂任务做完。把现有研究 [12, 13, 14, 17, 19] 归一归，常见的组织方式大致四类。"),

("上一个 Agent 的输出直接作为下一个 Agent 的输入",
 "串行（Sequential）：上一个 Agent 的输出直接喂给下一个，一个接一个往前推。语义依赖强的任务适合这么干，毛病也摆在那儿——任何一环卡住，整条链就停了。"),

("多个 Agent 各自独立给出方案，再通过投票或集成选出最终结果",
 "并行（Parallel）：几个 Agent 各自独立给方案，再投票或集成挑出最终结果。要“多视角投票”的任务最合适，Multi-Agent Debate [15] 就是一例。"),

("由一个总调度 Agent 将任务动态分派给下属",
 "分层（Hierarchical）：一个总调度 Agent 把任务动态派给下属，再把结果收回来整合。任务数量不固定、需要临时分派的场景适合它。"),

("对前一轮输出反复审查与修改，直至越过质量门槛",
 "反思（Reflexive）：同一个 Agent、或另设一个专门的 Critic Agent，对前一轮输出反复审、反复改，直到越过质量门槛（Reflexion [17]）。凡是要让质量收敛的任务，都得靠它。"),

("况且规划节点仅 3 个",
 "规划阶段本系统走串行。因为剧本、视觉、音乐这三类决策天生有先后：导演得先看到分镜大纲，才写得出镜头提示词；音乐制作人得先摸清导演的风格，才给得出匹配的音乐方案。"
 "何况规划节点就 3 个，硬套并行或分层只会平白多出协调开销。质量复检阶段则换成反思——让 Verifier 这个独立 Agent 给每个生成的镜头打分，不达标就退回上游重做。"
 "两样一搭，就是本系统“串行规划 + 反思复检”的整体协作模式。跟 ChatDev [14] 的瀑布式编排比，本系统特意把反思挪到生成阶段、而不是规划阶段："
 "文本规划这块让 LLM 高效串行过掉；真正费钱的“反思—重做”，专门留给方差最大的视觉生成环节。"),

("本系统的 State 共含 7 个核心字段",
 "接着说协作流转和上下文怎么传。LangGraph 靠一个共享的状态（State）对象在各节点之间递上下文。本系统的 State 一共 7 个核心字段，每个字段写入方、读取方都定得清清楚楚："),

("下游 Agent 读取 State 并非全盘接收",
 "下游 Agent 读 State 不是照单全收，而是按需挑字段。比如 Director 只读 character_bank 和 storyboard，不碰 user_intent——免得长上下文分散它写视觉指令的注意力；"
 "Verifier 则要同时读 character_bank、storyboard、prompt_pack，这三样凑齐才构成它打分的语义基线。这种字段级访问控制，相当于给每个 Agent 配了一张“只读视图”，"
 "既把角色隔开，又让上下文共享变得可控。单 Agent 范式做不到这点：它压根没有外部“视图”的概念，所有上下文全混在同一个 system prompt 和对话历史里，互相干扰。"),

("三类规划输出均采用严格的 JSON Schema 约束",
 "再说输出协议和字段级容错。三类规划输出都用严格的 JSON Schema 卡着：character_bank 必填 name / appearance / outfit / style，可选 accessories / signature_action；"
 "storyboard 必填 segment_id / label / start_time / end_time / description，可选 mood / characters / motion_hint；"
 "music_plan 必填 style / with_vocal / target_duration，可选 target_bpm / recommended_model / mood_curve。必填字段要是缺了，系统会自动在原任务描述后面加一句 JSON 错误提示，"
 "把同一个 Agent 重新调一遍（最多重试 2 次：第一次在提示词里加上“上次输出缺少 X 字段，请严格按 schema 重新生成”，第二次把温度降到 0.3 收紧约束）；还不行，就触发 Verifier 的回退路径。"
 "可选字段缺了，则填一个经验默认值，比如 mood 缺失时，按 description 里的关键词正则映射到 7 类基础情绪（happy / sad / energetic / calm / mysterious / romantic / epic）之一，"
 "characters 缺失时就继承上一个分镜的角色集合。这种“先严格校验、再分级容错”的协议，让整条规划链路对单个 LLM 的偶发失误挺扛造。"
 "开发后期统计了 200 多次规划运行，必填字段一次过的比例约 91%，加上自动重试后总通过率超过 99%。"),

("为何恰好是四个角色，而非三个或五个",
 "为什么偏偏是四个角色，不是三个或五个？本文的做法是把传统 MV 制作里“剧本—视觉—音乐—质检”这四个决策点，原样对应成四个 Agent，对应关系见表 3.1。"),

("本文做了两处合并",
 "跟已有工作的区别也得说一句。AutoGen [12] 是“对话式”多 Agent，MetaGPT [13] 拆的是“软件公司”角色，ChatDev [14] 用瀑布式编排。相比它们，本文做了两处合并。"
 "一处是把“摄影”并进 Director——在 AI 生成范式下，“摄影”说白了就是写视觉提示词，本就跟剪辑决策连得很紧；另一处是把“作曲”和“编曲”并进 Music Producer，"
 "一个 Agent 已经够给出统一的音乐方案了。只有 Verifier 单留着，因为质检逻辑和生成逻辑天生不是一回事，理应单独做成一个反思 Agent，这跟 Du 等人 [15] 的“多 Agent debate”思路也是相通的。"),

("Screenwriter（编剧 Agent）的角色定位为",
 "Screenwriter（编剧 Agent）的角色定位是“MV Screenwriter”，底层 LLM 用 GPT-4o，temperature 设 0.7（创意得适度发散），max_tokens 设 8192（够装下 6 镜的分镜加一份完整角色卡）。"
 "它的输入是 user_intent 和 music_analysis，输出是 character_bank 和 storyboard 两份 JSON。"),

("CharacterBank 在 MV 开头即固定全片所有角色",
 "关键设计 1 ── CharacterBank 一次性定义：CharacterBank 在 MV 一开头就把全片所有角色连同外貌、服装、配饰、风格标签一并定死，之后 Director 不再重新描述角色，"
 "只用角色 ID 引用——这是 3.5 节“跨镜头一致性”在语义层的根基。早期版本让 Director 自己描述每镜的角色，结果同一个“少女”在不同镜头里被写成不同的发色和服装，"
 "根子在于 Director 的上下文只有当前镜头的情节，看不到前后镜头。CharacterBank 把这种“重复描述的活儿”前移给 Screenwriter，从源头掐掉了跨镜头的描述漂移。"),

("每个 storyboard item 的 start_time / end_time 都必须严格对齐",
 "关键设计 2 ── 段落硬约束注入：每个 storyboard item 的 start_time / end_time 都必须严格对齐到 music_analysis.sections 给出的段落边界，"
 "副歌段落对 sing 镜头，主歌对 story 镜头，间奏给空镜或转场。Screenwriter 的 system prompt 里拿“Hard constraints”把这条规则明明白白列出来，"
 "而且分镜数量被卡在段落数 ± 1，免得 LLM 擅自合并或拆分段落。"),

("本文为每个分镜额外标注一个 mood",
 "关键设计 3 ── 情绪曲线保持一致：本文给每个分镜额外标一个 mood，从 7 类基础情绪里取；之后 Director 据此挑光线和构图模板，Music Producer 也据此调音乐生成模型的 prompt strength。"),

("Director（导演 Agent）的底层 LLM 同样选用",
 "Director（导演 Agent）底层 LLM 同样是 GPT-4o，temperature 设 0.5（视觉指令得稳），max_tokens 设 4096。输入是 character_bank 和 storyboard，"
 "要给每个分镜生成能直接执行的 image_prompt 和 video_prompt，并标上推荐的镜头景别和运镜方式。"),

("Director 不重新描述角色形象，而是直接将 character_bank",
 "关键策略 1 ── 自动注入角色锚点：Director 不重新描述角色长相，而是直接把 character_bank 里对应的字段，按“X is wearing Y, with Z hairstyle”这个固定句式拼到每条 prompt 的最前面。"
 "试了很多次，这种“按模板复述”对图像和视频模型保持角色一致帮助最大；与其让 LLM 自由发挥，不如这么换取更高的稳定性。"),

("sing 镜头默认输出 close-up shot",
 "关键策略 2 ── 按 label 分流视觉语言：sing 镜头默认输出 close-up shot, cinematic 3-point lighting, shallow depth of field, soft rim light；"
 "story 镜头默认输出 medium/long shot, handheld camera, naturalistic lighting, dolly-in motion。这套“视觉模板库”替 LLM 卸了视觉细节上的负担，让它能专心去写跟情节相关的特异性描述。"),

("Director 的输出始终是一条单独的 prompt 字符串",
 "关键策略 3 ── 提示词统一格式：系统虽然接了好几种不同的图像、视频模型，Director 的输出却始终是一条单独的 prompt 字符串，不为某个特定模型做特化，"
 "适配的活儿交给下游的 ShotRouter 和 Adapter 层。这样 Director 不用操心下游模型的细节，方法层和模型层也就解耦了。"),

("Music Producer（音乐制作人 Agent）的底层 LLM",
 "Music Producer（音乐制作人 Agent）底层 LLM 挑了轻一点的 Claude Haiku，temperature 设 0.4，max_tokens 设 1024（输出不过是个 dict）。"
 "它把 user_intent、music_analysis、storyboard.mood 综合一下，输出 music_plan，里面有 style（风格标签）、with_vocal（含不含人声）、target_bpm（目标 BPM）、"
 "target_duration（目标时长）、recommended_model（在 ACEStep / Suno / Lyria 里挑一个）、mood_curve（按段落给的情绪强度序列）。"),

("若用户已上传歌曲，Music Producer 便跳过选型步骤",
 "特殊路径 ── 上传音轨直通：用户要是已经传了歌，Music Producer 就跳过选型，只根据上传的音轨和 storyboard 给一条 mood_curve，给后面的视觉生成提供氛围引导；"
 "with_vocal、target_bpm 则改成直接从 music_analysis 读，省得再多调一次 LLM。"),

("在风格映射规则上",
 "风格映射规则上，用户意图里要是带了明确的风格关键词（像“国风”“赛博朋克”“复古迪斯科”这些），就锁死 style 字段，不再让 LLM 定；没带，才让 LLM 从 12 类常用风格里挑。"),

("Verifier（质检 Agent）的底层 LLM 选用 GPT-4o-mini",
 "Verifier（质检 Agent）底层 LLM 用 GPT-4o-mini（成本和速度优先），temperature 设 0.1（评分得稳），max_tokens 设 512。它的职能跟前三个错开：前三个负责生成，Verifier 只负责审。"
 "评分维度、阈值、重试策略详见 3.6 节。有个关键区别：Verifier 不拿任何上游 Agent 的“创造性输出”当参考答案，只以 character_bank 和原始 prompt 为基线，"
 "这样它的判断就独立在生成链路之外——要是让 Verifier 也掺进生成，它的评分就会出现自我合理化（self-justifying），外部审查也就白搭了。"),

("Planning Graph 以串行方式依次运行",
 "从执行顺序看，Planning Graph 串行地依次跑 Screenwriter → Director → Music Producer 三个生成节点，再把三者的输出汇进主 Pipeline、进入生成阶段；"
 "Verifier 节点不在规划链路上，而是每个镜头生成完之后才插进来的一个“复检分支”。串行是因为 Director 得先看到 storyboard 才写得出镜头提示词，"
 "Music Producer 得先摸到导演的风格（尤其是 storyboard.mood）才给得出匹配的 music_plan，中间是硬性的语义依赖。并行会把这依赖打散；分层结构则适合 Agent 多的场景，"
 "对只有 4 个 Agent 的本系统反倒是平添协调开销。"),

("下表给出三个规划 Agent 之间的",
 "接着说协作流转。下面这张表给出三个规划 Agent 之间的“握手节点”：上游写哪些字段、下游读哪些字段、两者之间有没有容错回退。"),

("框架据此构建一张有向无环图",
 "字段的流向都明确声明给了 LangGraph，框架据此搭出一张有向无环图（DAG），运行时校验“要读的字段必须已经被写过”。校验过不了"
 "（比如 Director 节点想读还没写入的 prompt_pack），框架就抛运行时异常，把重试控制器叫起来。"),

("再说明反思分支，即 Verifier 的反向跳转",
 "再说反思分支，也就是 Verifier 的反向跳转。Verifier 不是“挂在规划链路末尾”的终点节点，而是在每个生成阶段（图像生成 / 视频生成）完成后被异步触发的反向节点。"
 "它的评分通过 verifier_log 字段写回 State，再由重试控制器读：评分大于等于阈值就放行、进下一个镜头；小于阈值，就按下面的顺序逐级回退："),

("保持模型不变，将 cfg_scale 调至原值的 1.2 倍",
 "第 1 次失败：模型不动，把 cfg_scale 调到原值的 1.2 倍，换个随机种子重新生成（这是最快的一条路）；"),

("切换至 ShotRouter 中该镜头标签对应的 fallback 模型",
 "第 2 次失败：换成 ShotRouter 里该镜头标签对应的 fallback 模型（比如从 Seedance 切到 Veo 3.1）；"),

("回退至 Director，由其重新生成该镜头的 video_prompt",
 "第 3 次失败：触发“提示词重写分支”，退回 Director，让它重写这个镜头的 video_prompt，而且 Director 会把 Verifier 给的失败原因摘要附在提示词里"
 "（如 “Previous attempt was rejected for character outfit drift, please reinforce outfit description”）；"),

("以前一镜头的末帧作为占位帧，填满相同时长",
 "第 4 次失败（极少触发）：拿前一镜头的末帧当占位帧，填满相同时长，保证整条 Pipeline 不会因为某个单点卡死。"),

("verbal reinforcement",
 "这条反思分支跟 Reflexion [17] 讲的“verbal reinforcement”是一回事——把失败原因用自然语言反馈给上游 Agent，让 LLM 下一轮明确去改自己的输出，"
 "而不是把同一条提示词原样再跑一遍。在 30 条测试集上，开了这个分支后，单镜头通过率从 81% 提到了 95%，代价是平均时延多了约 15%。"),

("本研究在固定的 10 条测试集上，分别采用",
 "再说跟单 Agent 的对比。本研究在固定的 10 条测试集上，分别用 (i) 完整的四节点 LangGraph 规划图，和 (ii) 单个 GPT-4o Agent 一个人扛四项职责，各跑一遍生成。"
 "定性看下来：单 Agent 的输出大约 3/10 会出“Director 提示词里夹带音乐 BPM”这类角色漂移，多节点 LangGraph 里没出现；单 Agent 的 JSON 嵌套也比较乱，"
 "自动解析失败率约 1/10，多节点方案里各 Agent 各自吐独立的 JSON，解析失败率小于 1/30；CLIPScore 平均降了 0.039（−12.5%），综合 MOS 降 0.7 分（详见 7.3 节消融实验 B 组）。"),

("更深一层的区别在于",
 "再往深一层，“职责漂移”的根子在于：单 Agent 在长上下文里容易出现“刚写好的角色卡渗进正在写的音乐计划”这种情况，本质是 transformer 的 self-attention 没法在同一段上下文里分清"
 "“现在要做什么”和“之前做过什么”。多 Agent 则把“之前做过什么”用结构化字段（比如 character_bank、storyboard）明确递过去，而不是塞在自然语言的对话历史里，从根上躲开了职责混淆。"
 "这一点也跟 Hong 等人 [13] 在 MetaGPT 里得到的结论一致：用结构化输出而不是自由文本，是多 Agent 协作的关键工程做法。"),
]


def set_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


d = Document(PATH)
done = set()
for p in d.paragraphs:
    t = p.text
    for i, (snip, new) in enumerate(R):
        if i in done:
            continue
        if snip in t:
            set_text(p, new)
            done.add(i)
            break

missing = [(i, R[i][0][:24]) for i in range(len(R)) if i not in done]
if missing:
    print("[WARN] NOT matched:", missing)
else:
    d.save(PATH)
    print("saved", PATH, "| 共替换", len(done), "段")
